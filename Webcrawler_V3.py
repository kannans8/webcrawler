import streamlit as st
import pandas as pd
import requests
from bs4 import BeautifulSoup
from googlesearch import search
import re
import openpyxl
from collections import Counter
import heapq
import io
import google.generativeai as genai
import base64
import time
from docx import Document

import hashlib
import sqlite3
import difflib
from datetime import datetime
import os
import time
import pandas as pd

DB_PATH = "webcrawler_DB_11-2-2025.db"

def get_base64_of_bin_file(bin_file):
    with open(bin_file, "rb") as f:
        data = f.read()
    return base64.b64encode(data).decode()

st.set_page_config(page_title="Keyword Search AI", layout="wide")

base64_img = get_base64_of_bin_file("Picture1.png")

st.markdown(
    """
    <style>

        .stApp {
            background: url("data:image/jpg;base64,"""+base64_img+""" ") no-repeat center center fixed;
            background-size: cover;
            background-position: center 25px;
        }


        .title {
            position: absolute;
            top: 30%;               
            left: 50%;
            transform: translate(-50%, -50%);
            background: rgba(255, 255, 255, 0.9);
            padding: 20px 35px;
            border-radius: 15px;
            font-size: 42px;
            font-weight: bold;
            color: #003366;
            max-width: 500px;
            text-align: center;
            box-shadow: 0px 4px 12px rgba(0,0,0,0.2);
        }

    div[data-baseweb="input"] > input {
        width: 100px;      /* smaller width */
        height: 30px;      /* smaller height */
        font-size: 14px;   /* smaller text */
        padding: 4px 8px;
        border-radius: 8px;
        display: block;
        margin: 0 auto;    
    }


        .keyword-label, .pages-label {
            font-size: 60px;
            font-weight: bold;
            color: white;
            margin-top: 1px;
        }

 
        .logo-container {
            position: relative;
            top: 0px;
            left: 0px;
            z-index: 1;
        }
        .logo-container img {
            width: 150px;
        }



    }
    </style>





    <div class="title">Crawler AI_V2</div>
    <div style="position: fixed; bottom: 10px; left: 100px; color: white; font-size: 20px;">
        Gemini API used
    </div>
    """,
    unsafe_allow_html=True
)
# st.markdown('<div class="top-right-text"> </div>', unsafe_allow_html=True)
# st.markdown('<div class="top-right-text"> </div>', unsafe_allow_html=True)



# seconds = st.number_input("Set timer (in seconds):", min_value=1, max_value=3600, value=10, step=1)

# if st.button("Start Countdown"):
#     placeholder = st.empty()
#     for i in range(seconds, 0, -1):
#         placeholder.markdown(f"### ⏰ {i} seconds remaining")
#         time.sleep(1)
#     placeholder.markdown("## ✅ Time's up!")

def setup_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('''
        CREATE TABLE IF NOT EXISTS snapshots (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            Rank INTEGER,
            page_number INTEGER,
            Title TEXT,
            Link TEXT,
            Last_Updated TEXT,
            Summary TEXT,
            PDF_Links TEXT,
            url TEXT,
            date TEXT,
            html TEXT,
            hash TEXT,
            change_summary TEXT
        )
    ''')
    conn.commit()
    conn.close()
def get_last_snapshot(url):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('SELECT html, hash FROM snapshots WHERE url = ? ORDER BY id DESC LIMIT 1', (url,))
    row = cur.fetchone()
    conn.close()
    return row if row else (None, None)

import difflib

def compare_html(old_html, new_html):

    if not old_html:
        return "No previous snapshot available for comparison.", "Initial snapshot added"
    
    if not new_html:
        return "New HTML content is empty.", "Unable to compare — new HTML missing"

    old_lines = old_html.splitlines()
    new_lines = new_html.splitlines()
    
    diff = difflib.unified_diff(old_lines, new_lines, fromfile='previous', tofile='current', lineterm='')
    diff_list = list(diff)
    added = sum(1 for line in diff_list if line.startswith('+') and not line.startswith('+++'))
    removed = sum(1 for line in diff_list if line.startswith('-') and not line.startswith('---'))
    
    summary = f"Added {added} lines, removed {removed} lines" if added or removed else "Minor or no visible text changes"
    diff_text = "\n".join(diff_list)
    
    return diff_text if diff_text.strip() else "No visible text changes.", summary




def get_hash(content):
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

def save_snapshot(Rank, page_number, Title, Link, Last_updated, Summary, PDF_Links, url, html, hash_value, change_summary="Initial snapshot"):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('INSERT INTO snapshots (Rank, page_number, Title, Link, Last_updated, Summary, PDF_Links, url, date, html, hash, change_summary) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                (Rank, page_number, Title, Link, Last_updated, Summary, PDF_Links, url, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), html, hash_value, change_summary))
    conn.commit()
    conn.close()
    # print(f"💾 Snapshot saved for {url} ({len(html)} chars) — {change_summary}")



def get_last_updated(url):
    try:
        response = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(response.text, "html.parser")

        time_tag = soup.find("time")
        if time_tag and time_tag.get("datetime"):
            return time_tag["datetime"]
        elif time_tag:
            return time_tag.text.strip()

        meta_date = soup.find("meta", {"property": "article:modified_time"})
        if meta_date and meta_date.get("content"):
            return meta_date["content"]

        meta_date = soup.find("meta", {"name": "last-modified"})
        if meta_date and meta_date.get("content"):
            return meta_date["content"]

        match = re.search(r'\b\d{4}[-/]\d{2}[-/]\d{2}\b', soup.text)
        if match:
            return match.group(0)

    except Exception:
        return "N/A"

    return "-"

def summarize_text(text):


    # Primary_Keywords: ["external control arm", "external comparator", "historical control", "single-arm trial"]
    # Secondary_Keywords: ["guidance", "reflection paper", "draft", "guideline", "consultation", "regulatory decision"]
    # Concepts_to_include: ["Real-World Data (RWD)", "Real-World Evidence (RWE)", "evidence generation"]
    try:
        genai.configure(api_key="AIzaSyCWKOPXv28q0Tm39PNs0IlzPjFZpvXUb_A")
        model = genai.GenerativeModel('gemini-2.5-flash')
        prompt_instruction="You are an analystic , you are creating a report by searching a medical or government webistes to show any new updates are there in the website. Understand the line of texts or sentences and summarize the content in 200 words give importants to action items . make the content in bullet points . Exclude content related to advertisment, Contact Information. "
        full_prompt = f"{prompt_instruction}\n\n{text}"
        response = model.generate_content([full_prompt])
        return response.text
    except Exception :
        return "summarization error "


def summarize_text_full(text):
    try:
        genai.configure(api_key="AIzaSyCWKOPXv28q0Tm39PNs0IlzPjFZpvXUb_A")
        model = genai.GenerativeModel('gemini-2.5-flash')
        prompt_instruction="summary the excel information into a 2 page paragraph. give emphasis on summary column , on top give stats about how many links in the data and how many links had error and  give sitetation from the link column"
        full_prompt = f"{prompt_instruction}\n\n{text}"
        response = model.generate_content([full_prompt])
        return response.text
    except Exception as e:
        return "summarization error1 " + str(e)

# def summarize_text_full(text):
#     try:
#         prompt =  "summarize this content in 200 words and make bullet points : "
#         response = requests.post("http://127.0.0.1:8000/summarize", json={"text": text , "prompt": prompt})
#         if response.status_code == 200:
#             return response.json()["summary"]
#     except Exception as e:
#         return "summarization error1    -  " + str(e)
    
# def summarize_text(text):
#     try:
#         prompt =  "summarize this content in 200 words and make bullet points : "
#         response = requests.post("http://127.0.0.1:8000/summarize", json={"text":  text,  "prompt": prompt })
#         if response.status_code == 200:
#             return response.json()["summary"]
#     except Exception as e:
#         return "summarization error1 " + str(e)

def extract_text_from_url(url):
    try:
        response = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(response.text, "html.parser")
        for script in soup(["script", "style", "noscript"]):
            script.extract()
        text = " ".join(soup.stripped_strings)
        return text[:10000]
    except Exception:
        return ""

# def summarize_text(text, max_sentences=3):
#     if not text:
#         return "No content available."

#     sentences = re.split(r'(?<=[.!?]) +', text)
#     if len(sentences) <= max_sentences:
#         return " ".join(sentences)

#     words = re.findall(r'\w+', text.lower())
#     freq = Counter(words)

#     sentence_scores = {}
#     for sent in sentences:
#         sentence_words = re.findall(r'\w+', sent.lower())
#         score = sum(freq[w] for w in sentence_words if w in freq)
#         sentence_scores[sent] = score

#     best_sentences = heapq.nlargest(max_sentences, sentence_scores, key=sentence_scores.get)
#     return " ".join(best_sentences)


GOOGLE_SEARCH_API_KEY = "AIzaSyAN5Ke76-Z05fWBkajQlaYomPAOGPLgWbA"
# GOOGLE_SEARCH_API_KEY = 'AIzaSyDmbrTFpYfqQ-WuSdqZD62KsEINgGNDRc4'
GOOGLE_SEARCH_CX = "53438855d140e48fe"
def google_search_API(query, num=5):

    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "q": query,
        "key": GOOGLE_SEARCH_API_KEY,
        "cx": GOOGLE_SEARCH_CX,
        "num": num
    }
    resp = requests.get(url, params=params)
    resp.raise_for_status()
    data = resp.json()

    results = []
    if "items" in data:
        for item in data["items"]:
            results.append(item["link"])
    return results

# keyword = input("🔎 Enter your keyword: ").strip()

# print(f"\nSearching Google for: {keyword}\n")
# links = google_search(keyword, num=5)
# print("Top 5 links found:")
# for l in links:
#     print(" -", l)





def google_search_with_details_open(query, num_pages=1):
    results = []
    # all_urls = list(search(query, num_results=num_pages*5))
    all_urls = google_search_API(query, num=num_pages*5)

    for idx, url in enumerate(all_urls, start=1):
        page_number = (idx - 1) // 10 + 1
        rank = idx

        try:
            response = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
            soup = BeautifulSoup(response.text, "html.parser")
            title = soup.title.string.strip() if soup.title else "No Title"

            last_updated = get_last_updated(url)
            page_text = extract_text_from_url(url)
            summary = summarize_text(page_text)

            results.append([rank, page_number, title, url, last_updated, summary])
        except Exception as e:
            results.append([rank, page_number, "Error", url, "N/A", "error while load "+ str(e)])
    return results

def google_search_with_details_restried(query, num_pages=1):

    pharma_sites = [
        "ema.europa.eu",
        "raps.org",
        "ascpt.onlinelibrary.wiley.com"
    ]
    

    query = f"{query} site:{' OR site:'.join(pharma_sites)}"
    
    results = []
    all_urls = google_search_API(query, num=num_pages*5)

    for idx, url in enumerate(all_urls, start=1):
        page_number = (idx - 1) // 10 + 1
        rank = idx

        try:
            response = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
            soup = BeautifulSoup(response.text, "html.parser")
            title = soup.title.string.strip() if soup.title else "No Title"
            last_updated = ""

            # last_updated = get_last_updated(url)
            page_text = extract_text_from_url(url)
            summary = summarize_text(page_text)

            results.append([rank, page_number, title, url, last_updated, summary])
        except Exception:
            results.append([rank, page_number, "Error fetching title", url, "N/A", "Error to get prase the url"])
    return results

def google_search_with_predefined_url():
    results = []
    error_url = []

    all_urls = [
        'https://www.ema.europa.eu/en/development-reflection-paper-use-external-controls-evidence-generation-regulatory-decision-making-scientific-guideline',
        'https://www.ema.europa.eu/en/establishing-efficacy-based-single-arm-trials-submitted-pivotal-evidence-marketing-authorisation',
        'https://www.raps.org/news-and-articles/news-articles/2025/7/ema-proposes-reflection-paper-on-using-external-co',
        'https://ascpt.onlinelibrary.wiley.com/doi/10.1002/cpt.3684?af=R'
    ]

    for idx, url in enumerate(all_urls, start=1):
        page_number = (idx - 1) // 10 + 1
        rank = idx

        try:
            response = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
            soup = BeautifulSoup(response.text, "html.parser")
            title = soup.title.string.strip() if soup.title else "No Title"
            last_updated = get_last_updated(url)
            page_text = extract_text_from_url(url)
            summary = summarize_text(page_text)
            # summary = 'Testing'

            try:
                pdf = Search_pdf_links(url)
            except Exception:
                pdf = ['No PDF found']

            print(f"\n🔍 Checking: {url}")
            html = response.text
            current_hash = get_hash(html)
            old_html, old_hash = get_last_snapshot(url)

            if old_hash is None:
                print("🆕 First snapshot — saving initial version.")
                save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf),
                              url , html, current_hash, "Initial snapshot")
            elif current_hash == old_hash:
                print("✅ No changes detected.")
                save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf),
                              url ,html, current_hash, "No changes detected")
            else:
                try:
                    print("⚠️ Changes detected!")
                    diff_text, summary1 = compare_html(old_html, html)
                    save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf),
                                  url , html, current_hash, summary1)
                    print(f"📦 Snapshot updated with summary: {summary1}")
                except Exception as e:
                    summary1 = "Error comparing HTML"
                    print("Error comparing HTML:", str(e))
                    save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf),
                                 url, html, current_hash, summary1)

            results.append([rank, page_number, title, url, last_updated, summary, str(pdf)])

        except Exception as e:
            results.append([rank, page_number, "Error fetching title", url, "N/A", "Error parsing URL", "na"])
            error_url.append(url)
            print("❌ Error in predefined URL:", str(e))

    return results, error_url







# def google_search_with_predefined_url():
#     results = []
#     error_url = []


#     all_urls = ['https://www.ema.europa.eu/en/development-reflection-paper-use-external-controls-evidence-generation-regulatory-decision-making-scientific-guideline',
#                 'https://www.ema.europa.eu/en/establishing-efficacy-based-single-arm-trials-submitted-pivotal-evidence-marketing-authorisation',
#                 'https://www.raps.org/news-and-articles/news-articles/2025/7/ema-proposes-reflection-paper-on-using-external-co',
#                 'https://ascpt.onlinelibrary.wiley.com/doi/10.1002/cpt.3684?af=R']
    

#     for idx, url in enumerate(all_urls, start=1):
#         page_number = (idx - 1) // 10 + 1
#         rank = idx


#         try:
#             response = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
#             soup = BeautifulSoup(response.text, "html.parser")
#             title = soup.title.string.strip() if soup.title else "No Title"
#             last_updated = get_last_updated(url)
#             # page_text = extract_text_from_url(url)
#             # summary = summarize_text(page_text)
#             summary = 'Testing'
#             try:
#                 pdf = Search_pdf_links(url)
#             except Exception:
#                 pdf = ['No PDF found']
            

#             print(f"\n🔍 Checking: {url}")
#             html = response.text
#             print("1")

#             current_hash = get_hash(html)
#             print("2")
#             old_html, old_hash = get_last_snapshot(url)
#             print("3")


#             if old_hash is None:
#                 print("🆕 First snapshot — saving initial version.")
#                 save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf), url, html, current_hash, "Initial snapshot")

#             if current_hash == old_hash:
#                 print("✅ No changes detected.")
#                 save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf), url, html, current_hash, "No changes detected")
#             else:
#                 try:
#                     print("⚠️ Changes detected!")
#                     diff_text, summary1 = compare_html(old_html, html)
#                     save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf), url, html, current_hash, summary1)
#                     print(f"Snapshot updated with summary: {summary1}")
#                 except Exception as e:
#                     summary1 = "error comparing HTML" 
#                     print("Error comparing HTML:", str(e))
#                     save_snapshot(rank, page_number, title, url, last_updated, summary, str(pdf), url, html, current_hash, summary1)

#             results.append([rank, page_number, title, url, last_updated, summary, str(pdf)])
#         except Exception as e:
#             results.append([rank, page_number, "Error fetching title", url, "N/A", "Error to get prase the url", "na"])
#             error_url.append(url)
#             print("error in predefined url " + str(e))
#     return results , error_url

def Search_pdf_links(page_url):
    page = requests.get(page_url)
    page.raise_for_status()
    soup = BeautifulSoup(page.text, "html.parser")
    pdf_links = [a["href"] for a in soup.find_all("a", href=True) if a["href"].endswith(".pdf")]
    full_pdf_links = []
    for link in pdf_links:
        if not link.startswith("http"):
            from urllib.parse import urljoin
            link = urljoin(page_url, link)
        full_pdf_links.append(link)
        print(f"Downloading PDF: {link}")
    return full_pdf_links

def save_to_excel(keyword, data):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = keyword[:31]
    sheet.append(["Rank", "Page Number", "Title", "Link", "Last Updated", "Summary", "PDF Links"])

    for row in data:
        sheet.append(row)

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def save_to_word(keyword, data):
    doc = Document()
    doc.add_heading(f"Search Results for '{keyword}'", level=1)
    doc.add_paragraph(
    f"This document contains the summarized search results for the keyword "
    f"'{keyword}'. " + str(data)
)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def get_snapshots_by_filter(selected_filters):
    conn = sqlite3.connect(DB_PATH)
    query = """
        SELECT Rank, page_number, Title, Link, Last_Updated, Summary, PDF_Links, change_summary, date
        FROM snapshots
        WHERE 1=1
    """

    conditions = []
    if "Change Detected" in selected_filters:
        conditions.append("(change_summary LIKE '%Added%' OR change_summary LIKE '%removed%' OR change_summary LIKE '%Changes detected%')")
    if "No Change Detected" in selected_filters:
        conditions.append("change_summary = 'No changes detected'")
    if "New Run" in selected_filters:
        conditions.append("change_summary = 'Initial snapshot'")

    if conditions:
        query += " AND (" + " OR ".join(conditions) + ")"

    query += " ORDER BY id DESC"

    df = pd.read_sql_query(query, conn)
    conn.close()
    return df





st.markdown('<p class="keyword-label">Enter the Keyword:</p>', unsafe_allow_html=True)
# st.markdown('<div class="small-input"></div>', unsafe_allow_html=True)
keyword = st.text_input("", key="keyword_input", placeholder="Type here...")

# Pages Input
st.markdown('<p class="pages-label">Number of Pages to Search:</p>', unsafe_allow_html=True)
# pages = st.number_input("", min_value=1, max_value=100, value=1, step=1)
pages = st.slider("", 1, 100, 1)

# search_type = st.radio(
#     "Choose Search Type:",
#     ("Restricted (Pharma Sites)", "Open Search", "Predefined url"),
#     horizontal=True
# )
search_type = st.radio(
    "Choose Search Type:",
    ("Open Search", "Predefined url"),
    horizontal=True
)


# Initialize session state for results
if "search_results" not in st.session_state:
    st.session_state.search_results = []
if "full_summary" not in st.session_state:
    st.session_state.full_summary = ""

# Run Search button
if st.button("Run Search"):
    if keyword.strip():
        with st.spinner("Searching and summarizing... Please wait for 2 mins "):
            
            # Determine search type and run search
            if search_type == "Restricted (Pharma Sites)":
                # Replace with actual search function
                search_results = google_search_with_details_restried(keyword, num_pages=pages)
                st.session_state.search_results  = search_results
            elif search_type == "Predefined url":
                print("predefined url selected")
                # Replace with actual search function
                search_results, error_url = google_search_with_predefined_url()
                # print(error_url)
                search_results = [ result for result in search_results if result[3] not in error_url]
                st.session_state.search_results  = search_results

            else:
                search_results = google_search_with_details_open(keyword, num_pages=pages)
                st.session_state.search_results  = search_results
            

            st.session_state.full_summary = summarize_text_full(search_results)
            # st.session_state.full_summary = 'Testing full summary'

            # Prepare files for download
            word_file = save_to_word(keyword, st.session_state.full_summary)
            excel_file = save_to_excel(keyword, st.session_state.search_results)

        st.success("Search completed") 

        # Save files in session state for download buttons
        st.session_state.word_file = word_file
        st.session_state.excel_file = excel_file


if st.session_state.search_results:
    col1, col2 = st.columns([2,1])

    with col1:
        st.write("Preview Table")
        if search_type == "Predefined url":
            df = pd.DataFrame(
                st.session_state.search_results,
                columns=["Rank", "Page Number", "Title", "Link", "Last Updated", "Summary", "PDF Links"]
            )
        else:
            df = pd.DataFrame(
                st.session_state.search_results,
                columns=["Rank", "Page Number", "Title", "Link", "Last Updated", "Summary"]
            )            
        styled_df = df.style.set_properties(**{
            'background-color': "#f2f3f4",
            'color': '#000080',
            'border-color': '#000080'
        }).set_table_styles([
            {'selector': 'th', 'props': [('background-color', "#170E7A"),
                                         ('color', "#4D741E"),
                                         ('font-weight', 'bold')]}
        ])
        st.dataframe(styled_df, use_container_width=True)

    with col2:
        st.write("Full Summary")
        st.markdown(
        f"""
        <div style="
            color: #000080; 
            font-weight: bold; 
            background-color: #FFFFFF; 
            padding: 10px; 
            border-radius: 8px;
            height: 400px; 
            overflow-y: auto;
            white-space: pre-wrap;
        ">
            {st.session_state.full_summary}
        </div>
        """,
        unsafe_allow_html=True
    )


# Download buttons if available
if "word_file" in st.session_state and "excel_file" in st.session_state:
    st.download_button(
        label="📥 Download Fully Summary (Word)",
        data=st.session_state.word_file,
        file_name=f"{keyword}_results.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.download_button(
        label="📥 Download Results as Excel",
        data=st.session_state.excel_file,
        file_name=f"{keyword}_results.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

st.markdown("## Snapshot History Viewer")
st.write("Use the filters below to view snapshots by change type:")

colf1, colf2, colf3 = st.columns(3)
with colf1:
    change_detected = st.checkbox("Change Detected", value=True)
with colf2:
    no_change = st.checkbox("No Change Detected", value=False)
with colf3:
    new_run = st.checkbox("New Run", value=False)

selected_filters = []
if change_detected:
    selected_filters.append("Change Detected")
if no_change:
    selected_filters.append("No Change Detected")
if new_run:
    selected_filters.append("New Run")

# Fetch filtered data from DB
snapshot_df = get_snapshots_by_filter(selected_filters)

st.divider()

if not snapshot_df.empty:
    st.dataframe(
        snapshot_df.style.set_properties(**{
            'background-color': "#eef5ff",
            'color': '#001f4d',
            'border-color': '#001f4d'
        }).set_table_styles([
            {'selector': 'th', 'props': [('background-color', "#003366"),
                                         ('color', "white"),
                                         ('font-weight', 'bold')]}
        ]),
        use_container_width=True,
        height=400
    )
else:
    st.info("No records found for the selected filters.")

if st.button("Refresh Snapshot Data"):
    st.rerun()


